from docx import Document
import os
import win32com.client as win32
import re
from datetime import datetime
import pypandoc
import requests
from docx.shared import Pt
import locale
import time
import pyautogui
import tkinter as tk
from tkinter import ttk
from PIL import Image, ImageTk
import threading
import pythoncom
import calendar
import pywintypes
import traceback

USERNAME = os.getenv("USERNAME")
# Definir o local para o formato brasileiro
locale.setlocale(locale.LC_TIME, 'pt_BR.UTF-8')

# Caminhos dos arquivos
pasta_dados = fr"C:\Users\{USERNAME}\Documents\empresas"
template_file_path = fr"C:\Users\{USERNAME}\Documents\template\2024 - LTCAT PADRÃO - MODELO NR 20 - DR FERNANDO.doc"
output_pdf_path = fr"C:\Users\{USERNAME}\Desktop\ltcat"
doc_reorganizado_path = fr"C:\Users\{USERNAME}\Desktop\ltcat\teste.docx"

# Obter a data de hoje
hoje = datetime.now()

ano_atual = datetime.now().year

mes_atual = datetime.now().month

def convert_to_docx(arquivo):
    try:
        # Abrir o Word
        word = win32.Dispatch("Word.Application")
        word.Visible = False 
        doc = word.Documents.Open(arquivo)

        # Salvar como .docx
        if(arquivo.split('.')[1] == 'rtf'):
            output_arquivo = arquivo.replace('.rtf', '.docx')
        else:
            output_arquivo = arquivo.replace('.doc', '.docx')
            
        doc.SaveAs(output_arquivo, FileFormat=16) 
        doc.Close()
        return output_arquivo

    except Exception as e:
        print(f"Erro ao processar o arquivo: {e}")

def extrair_nome_documento(doc):
    for para in doc[1].paragraphs:
        # Verifica se a linha anterior contém os cabeçalhos relevantes
        if "POSTOS YANI LTDA" in para.text:
            texto_separado = para.text.split(':')
            return texto_separado[1].strip()
    return None

def obter_nome_documento(file_path):
    # Extrai o nome do arquivo (sem o caminho completo)
    file_name = os.path.basename(file_path)
    
    # Expressão regular para capturar o ano e o mês no formato esperado
    match = re.search(r"(\d{4})\s*-\s*(\w+)", file_name)
    
    if match:
        year = match.group(1)
        month = match.group(2).upper()  # Converter para maiúsculas
        
        # Dicionário para mapear os meses para o formato desejado
        meses = {
            "JANEIRO": "JANEIRO",
            "FEVEREIRO": "FEVEREIRO",
            "MARÇO": "MARÇO",
            "ABRIL": "ABRIL",
            "MAIO": "MAIO",
            "JUNHO": "JUNHO",
            "JULHO": "JULHO",
            "AGOSTO": "AGOSTO",
            "SETEMBRO": "SETEMBRO",
            "OUTUBRO": "OUTUBRO",
            "NOVEMBRO": "NOVEMBRO",
            "DEZEMBRO": "DEZEMBRO"
        }
        
        # Verifica se o mês está no dicionário e retorna o formato desejado
        if month in meses:
            return f"{meses[month]} DE {year}"
        else:
            raise ValueError("Mês não reconhecido no nome do arquivo.")
    else:
        raise ValueError("Formato do nome do arquivo inválido. Não foi possível extrair ano e mês.")

def obter_data_hoje_formatacao_documento():
    # Obtém a data atual
    data_atual = datetime.now()
    # Formata a data como 'DD-MM-YYYY'
    data_formatada = data_atual.strftime("%d-%m-%Y")
    return data_formatada

def read_word_file(file_path):
    doc = Document(file_path)
    content = []
    for para in doc.paragraphs:
        content.append(para.text)
    return content, doc

def converter_rtf_para_docx(input_path, output_path):
    pypandoc.convert_file(input_path, 'docx', outputfile=output_path)
    
def obter_cnpj_e_data(output_docx_path):
    # Lê o documento Word (.docx)
    doc = Document(output_docx_path)
    cnpj_pattern = r"\d{2}\.\d{3}\.\d{3}/\d{4}-\d{2}"
    data_pattern = r"\d{2}/\d{2}/\d{4}"

    cnpj = None
    data = None

    # Primeira tentativa: busca o CNPJ em todas as células de todas as tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if not cnpj:
                    cnpj_match = re.search(cnpj_pattern, cell.text)
                    if cnpj_match:
                        cnpj = cnpj_match.group()
                        break
            if cnpj:
                break
        if cnpj:
            break

    # Segunda tentativa: buscar a data apenas na tabela que contém "RUÍDO CONTINUO"
    for table in doc.tables:
        table_text = " ".join(cell.text for row in table.rows for cell in row.cells)
        if "RUÍDO CONTINUO" in table_text.upper():
            for row in table.rows:
                for cell in row.cells:
                    if not data:
                        data_match = re.search(data_pattern, cell.text)
                        if data_match:
                            data = data_match.group()
                            break
                if data:
                    break
            break

    return cnpj, data

def consulta_cartao_cnpj(cnpj):
    # URL e cabeçalhos da API
    cnpj_consulta = cnpj.replace('.', '').replace('/', '').replace('-', '')
    url = f"https://api.cnpja.com/office/{cnpj_consulta}"
    headers = {
        "Authorization": "ec1ea1b9-cb4f-460d-8ac1-3cba089fb252-1b1de35e-1616-46e6-9f76-737d7a18194d"
    }
    
    # Enviar a solicitação GET
    response = requests.get(url, headers=headers)

    if response.status_code == 200:
        data = response.json()
        
        # Obtenção dos valores com mensagens padrão caso não encontrados
        cnpj = data.get("taxId", "Não informado")
        company = data.get("company", {})
        nome_empresa = company.get("name", "Não informado").upper()
        porte = company.get("size", {}).get("acronym", "Não informado")
        nome_fantasia = data.get("alias", "Não informado").upper()
        data_abertura = data.get("founded", "Não informado").upper()
        data_sit_cad = data.get("statusDate", "Não informado")
        status = data.get("status", {})
        status_text = status.get('text', 'Não informado').upper()

        # Natureza jurídica
        nature = company.get('nature', {})
        nature_id = nature.get('id', 'Não informado')
        nature_text = nature.get('text', 'Não informado').upper()
        codigo_desc = f"{str(nature_id)[:3]}-{str(nature_id)[-1]} - {nature_text}"

        # Endereço
        address = data.get("address", {})
        logradouro = address.get('street', 'Não informado').upper()
        numero = address.get('number', 'Não informado')
        complemento = (address.get('details') or 'Não informado').upper()
        bairro = address.get('district', 'Não informado').upper()
        municipio = address.get('city', 'Não informado').upper()
        uf = address.get('state', 'Não informado').upper()
        cep = address.get('zip', 'Não informado')

        # Telefones
        phones = data.get("phones", [])
        phone_list = [
            f"({telefone.get('area', 'Área não informada')}) {telefone.get('number', 'Número não informado')}" for telefone in phones
        ] if phones else ["Não informado"]

        # Emails
        emails = data.get("emails", [])
        email_list = [
            email.get('address', 'Não informado').upper() for email in emails
        ] if emails else ["Não informado"]

        # Atividade principal
        main_activity = data.get("mainActivity", {})
        codigo = main_activity.get('id', 'Não informado')
        if codigo != 'Não informado':
            codigo = f"{str(codigo)[:2]}.{str(codigo)[2:4]}-{str(codigo)[4:5]}-{str(codigo)[5:]}"
        atividade = main_activity.get('text', 'Não informado').upper()
        codigo_completo = f"{codigo} - {atividade}"

        # Atividades secundárias
        atividade_sec = data.get("sideActivities", [])
        atividade_sec_text = ', '.join(item['text'] for item in atividade_sec) if atividade_sec else "Não informada"

        # Retorno dos dados coletados ou valores padrão
        return {
            "cnpj": cnpj,
            "nome_empresa": nome_empresa,
            "porte": porte,
            "nome_fantasia": nome_fantasia,
            "data_abertura": data_abertura,
            "data_sit_cad": data_sit_cad,
            "status_text": status_text,
            "codigo_desc": codigo_desc,
            "logradouro": logradouro,
            "numero": numero,
            "complemento": complemento,
            "bairro": bairro,
            "municipio": municipio,
            "uf": uf,
            "cep": cep,
            "telefones": phone_list,
            "emails": email_list,
            "codigo_completo": codigo_completo,
            "atividade_sec_text": atividade_sec_text
        }
    else:
        return {"error": f"Erro ao consultar API: Código {response.status_code}"}

def format_date(date_str):
    try:
        date_obj = datetime.fromisoformat(date_str.replace('Z', '+00:00'))
        return date_obj.strftime('%d/%m/%Y')
    except ValueError:
        return 'Data inválida'
            
def criar_novo_run(paragrafo, texto, negrito=False, fonte="Verdana", tamanho=8):
    # Cria o novo run com o texto
    novo_run = paragrafo.add_run(texto)
    
    # Define o negrito explicitamente de várias formas
    novo_run.bold = negrito
    novo_run.font.bold = negrito
    
    # Define a fonte e o tamanho
    novo_run.font.name = fonte
    novo_run.font.size = Pt(tamanho)
    
    return novo_run
        
def formatar_data_tabela(doc, replacements):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Verifica se o texto da célula corresponde ao formato de data '00.06.2023'
                if re.match(r'\d{2}\.\d{2}\.\d{4}', cell.text):
                    data_original = cell.text
                    # Substitui o texto da célula pelo valor correspondente no dicionário replacements
                    if data_original in replacements:
                        data_substituicao = replacements[data_original]
                        cell.text = data_substituicao.replace('-','/')
   
def selecionando_conteudo_setor_adm(word_doc_path, start_line_text):
    word = win32.Dispatch("Word.Application")
    word.Visible = True  
    doc = word.Documents.Open(word_doc_path)

    time.sleep(3) 

    found = False
    for paragraph in doc.Paragraphs:
        if start_line_text in paragraph.Range.Text:
            found = True
            start_range = paragraph.Range
            end_range = doc.Range(Start=start_range.End, End=doc.Content.End)

            # Selecionar o conteúdo
            start_range.Select()
            pyautogui.click()
            time.sleep(1)
            end_range = doc.Range(Start=start_range.End, End=doc.Content.End)
            end_range.Select()
            
            time.sleep(2) 

            pyautogui.hotkey('ctrl', 'c')
            time.sleep(8)  
            print("Conteúdo copiado com sucesso.")
            break

    if not found:
        print(f"Texto '{start_line_text}' não encontrado no documento.")

    doc.Close()
    word.Quit()
    
def colar_conteudo_em_pag_15(destination_path, progress_label):
    word = win32.Dispatch("Word.Application")
    word.Visible = True
    doc = word.Documents.Open(destination_path)

    time.sleep(3)

    try:
        # Navegar até a página especificada
        selection = word.Selection
        selection.GoTo(What=1, Which=1, Count=15)
        time.sleep(2)

        # Mover o cursor para baixo para pular a linha desejada
        selection.MoveDown(Unit=5, Count=1)
        selection.TypeParagraph() # Adicionar uma linha em branco
        time.sleep(2)
        pyautogui.hotkey('ctrl', 'v')
        time.sleep(2)
        print("Conteúdo colado com sucesso.")
        
    except Exception as e:
        print(f"Erro ao tentar colar na página {15}: {e}")

    # Salvar e fechar o documento
    base_name = os.path.basename(destination_path)  
    folder_path = os.path.dirname(destination_path)

    # Separar o nome do arquivo e a extensão
    name, extension = os.path.splitext(base_name) 

    # Adicionar " Editado" ao nome do arquivo
    new_name = f"{name} Editado{extension}"

    # Juntar o novo nome com o caminho da pasta
    new_destination_path = os.path.join(folder_path, new_name)
    
    doc.SaveAs(new_destination_path)
    doc.Close()
    word.Quit()
    time.sleep(2)
    
    if excluir_tabelas_formatar_e_reorganizar_documento(new_destination_path, word, progress_label):
        word = win32.Dispatch("Word.Application")
        word.Visible = False
        template_doc = word.Documents.Open(new_destination_path)
        template_range = template_doc.Content

        placeholder = "{{tabela}}"
        find_placeholder = template_range.Find
        find_placeholder.Text = placeholder

        if find_placeholder.Execute():
            print(f"Marcador {placeholder} encontrado. Colando a tabela.")
            find_placeholder.Parent.Paste()  # Colar o conteúdo no local do marcador

            # Seleciona a última tabela no documento
            last_table = template_doc.Tables(template_doc.Tables.Count)
            last_table_range = last_table.Range
            last_table_range.Collapse(0)  # Colapsa o cursor para o final da tabela

            # Remover parágrafos vazios após a última tabela
            remove_blank_paragraphs_after_table(template_doc, last_table_range)
        else:
            print(f"Marcador {placeholder} não encontrado no documento {destination_path}.")

        template_doc.SaveAs(new_destination_path)
        template_doc.Close()
        progress_label.config(text="Finalizado formatação das tabelas copiadas")
        progress_label.config(text=f"Documento final salvo em: {new_destination_path}")
        print(f"Documento final salvo em: {new_destination_path}")
    else:
        print(f"Erro ao reorganizar o documento {new_destination_path}")
        progress_label.config(text=f"Erro ao reorganizar o documento {new_destination_path}")

    return new_destination_path

def remove_blank_paragraphs_after_table(doc, table_range):
    end_of_table = table_range.End
    range_after_table = doc.Range(end_of_table, doc.Content.End)

    while range_after_table.Text.strip() == '':
        range_after_table.Delete()  
        range_after_table = doc.Range(end_of_table, doc.Content.End)  

    if range_after_table.Paragraphs.Count > 0:
        for i in range(range_after_table.Paragraphs.Count, 0, -1): 
            para = range_after_table.Paragraphs(i)
            if para.Range.Text.strip() == '':
                para.Range.Delete()  
                               
def excluir_tabelas_formatar_e_reorganizar_documento(doc_path, word, progress_label):
        progress_label.config(text="Realizando formatação das tabelas copiadas")
        try:
            word = win32.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(doc_path)

            # Remover tabelas com linhas em branco
            tables_to_remove = []
            for i in range(1, doc.Tables.Count + 1):
                table = doc.Tables(i)
                blank_row_found = False

                # Verifica cada célula da tabela para identificar linhas em branco
                for row in table.Rows:
                    row_text = ''.join([cell.Range.Text.strip() for cell in row.Cells])
                    if not row_text:  # Se o texto da linha for vazio
                        blank_row_found = True
                        break  # Não precisa verificar mais células desta tabela

                # Se encontrou linha em branco, marca a tabela para exclusão
                if blank_row_found:
                    tables_to_remove.append(i)

            # Remove as tabelas marcadas
            for index in reversed(tables_to_remove):  # Reverte a ordem para remover sem perder a referência
                doc.Tables(index).Delete()

            # Formatar todas as células das tabelas restantes
            for table in doc.Tables:
                for row in table.Rows:
                    for cell in row.Cells:
                        cell_range = cell.Range
                        cell_range.Font.Name = 'Verdana'
                        cell_range.Font.Size = 8

            # Reorganizar o documento
            range_total = doc.Range()
            paragraphs = range_total.Paragraphs
            setor_encontrado = False
            setor_atual = ""
            cargo_atual = ""
            first_setor = True  # Indica se é o primeiro setor encontrado
            cargo_encontrado = False  # Indica se já encontramos o cargo após o setor

            for paragraph in paragraphs:
                texto = paragraph.Range.Text.strip()

                # Verifica se o parágrafo começa com "Setor:"
                if texto.startswith("Setor:"):
                    # Quebra de página antes de cada novo setor (exceto o primeiro)
                    if not first_setor and texto != setor_atual:
                        paragraph.Range.InsertBreak(7)  # 7 é o valor de wdPageBreak (quebra de página)
                    
                    # Atualiza o setor atual e define que o próximo setor não é o primeiro
                    setor_atual = texto
                    first_setor = False
                    cargo_encontrado = False  # Reseta o cargo encontrado para o novo setor
                    continue  # Continua para não adicionar quebras de página entre "Setor:" e o primeiro "Cargo:"

                # Verifica se o parágrafo começa com "Cargo:"
                if texto.startswith("Cargo:"):
                    if not cargo_encontrado:
                        # Se for o primeiro cargo do setor, não insere a quebra de página
                        cargo_encontrado = True
                    else:
                        # Quebra de página antes de cada cargo após o primeiro
                        paragraph.Range.InsertBreak(7)

            # Seleciona todo o conteúdo reorganizado e copia para a área de transferência
            range_total.Select()
            word.Selection.Copy()  # Copia o conteúdo reorganizado para a área de transferência
            doc.Close()  # Fecha o documento
            return True
        except pywintypes.com_error as e:
            progress_label.config(text="Erro ao excluir tabelas, formatar e reorganizar o documento: {e}")
            print(f"Erro ao excluir tabelas, formatar e reorganizar o documento: {e}")
            traceback.print_exc()
            return False
                         
def substituir_texto_no_documento(doc, replacements, caminho_final, nome_documento, data_documento):
    formatar_data_tabela(doc, replacements)
    ja_preencheu = False
    ja_preencheu_data = False
    
    def substituir_em_runs(paragrafo, runs, chave, valor):
        full_text = ''.join([run.text for run in runs])

        if chave in full_text:
            # Substitui a chave pelo valor mantendo o resto do texto
            novo_texto = full_text.replace(chave, valor)

            # Remove o texto dos runs existentes
            for run in runs:
                # Remover cor de destaque e definir cor do texto como preto
                run.font.color.rgb = None  # Reseta a cor do texto
                if run._element.xpath('.//w:highlight'):
                    run._element.remove(run._element.xpath('.//w:highlight')[0])  # Remove o destaque de cor

                run.text = ''  # Limpa o conteúdo do run

            # Recria os runs com o novo texto e aplica negrito ao valor
            partes = novo_texto.split(valor)
            if len(partes) == 2:
                # Primeiro run com o texto antes do valor
                criar_novo_run(paragrafo, partes[0], fonte="Verdana", tamanho=8)

                # Novo run para o valor com negrito
                criar_novo_run(paragrafo, valor, negrito=True, fonte="Verdana", tamanho=8)

                # Run com o texto restante
                criar_novo_run(paragrafo, partes[1], fonte="Verdana", tamanho=8)

    # Substituição nos parágrafos
    for p in doc.paragraphs:
        for chave, valor in replacements.items():
            if f"{{{{{chave}}}}}" in p.text:
                substituir_em_runs(p, p.runs, f"{{{{{chave}}}}}", valor)

    # Substituição nas tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for chave, valor in replacements.items():
                        if f"{{{{{chave}}}}}" in p.text:
                            substituir_em_runs(p, p.runs, f"{{{{{chave}}}}}", valor)

    # Substituição direta de texto em parágrafos
    for para in doc.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in para.text:
                # Guarda o texto original e limpa os runs atuais
                texto_antigo = para.text
                for run in para.runs:
                    run.text = ''
                
                # Cria novos runs para cada parte do texto substituído
                partes = texto_antigo.split(old_text)
                                        
                for i, parte in enumerate(partes):
                    print(f"Passou aqui da parte: {parte}")
                    criar_novo_run(para, parte, fonte="Verdana", tamanho=8)
                    if i < len(partes) - 1:
                        print(f"Passou aqui do novo texto: {new_text}")
                        tamanho_fonte = 8
                        
                        if(new_text == nome_documento and ja_preencheu == False):
                            ja_preencheu = True
                            tamanho_fonte = 18
                        
                        if(new_text == data_documento and ja_preencheu_data == False):
                            ja_preencheu_data = True
                            tamanho_fonte = 12
                            
                        criar_novo_run(para, new_text, negrito=True, fonte="Verdana", tamanho=tamanho_fonte)

    doc.save(caminho_final)
    print('Salvando documento alterado...')

def save_as_pdf(doc_path, output_pdf_path):
    word = win32.Dispatch("Word.Application")
    word.Visible = False
    doc = word.Documents.Open(doc_path)
    doc.SaveAs(output_pdf_path, FileFormat=17)  # Formato PDF
    doc.Close()
    word.Quit()

def converter_data_pt_br(data):
    data_obj = datetime.strptime(data, '%Y-%m-%d')

    return data_obj.strftime('%d/%m/%Y')

def processar_arquivos(progress_label, progress_bar):
    pythoncom.CoInitialize()
    progress_label.config(text="Iniciando processos")
    time.sleep(1)
    
    arquivos_dados = [f for f in os.listdir(pasta_dados) if f.endswith('.rtf')]

    for arquivo in arquivos_dados:
        progress_label.config(text=f"Processando arquivo: {arquivo}...")
        time.sleep(1)
    
        progress_label.config(text="Convertendo os arquivos para DOCX")
        output_docx_path = convert_to_docx(pasta_dados+'\\'+arquivo)
        time.sleep(3)
        template_output_file_path = convert_to_docx(template_file_path)

        progress_label.config(text="Leitura do arquivo da Empresa")
        original_doc = read_word_file(output_docx_path)
        
        progress_label.config(text="Extraindo nome e data do documento")
        nome_documento = extrair_nome_documento(original_doc)
        data_documento = obter_nome_documento(output_docx_path)
        data_formatacao_documento = obter_data_hoje_formatacao_documento()
        
        progress_label.config(text="Obtendo CNPJ do documento")
        dados_doc_empresa = obter_cnpj_e_data(output_docx_path)
        cnpj = dados_doc_empresa[0]
        data_documento_empresa = dados_doc_empresa[1]
        
        progress_label.config(text="Obtendo dados sobre o CNPJ através da API")
        infos_cartao_cnpj = consulta_cartao_cnpj(cnpj)
        
        progress_label.config(text="Selecionando informações a partir do Setor: ADMINISTRATIVO")
        selecionando_conteudo_setor_adm(output_docx_path,"Setor: ADMINISTRATIVO")
        time.sleep(5)
        pyautogui.hotkey('enter')
        time.sleep(2)
        
        progress_label.config(text="Realizando colagem do conteúdo no DESCRIÇÃO DAS ATIVIDADES E DOS RISCOS AMBIENTAIS")
        template_editado = colar_conteudo_em_pag_15(template_output_file_path, progress_label)

        #Convertendo formado da data de hoje..
        hoje = datetime.now()
        ano_atual = datetime.now().year
        mes_atual = datetime.now().month

        # Obter o nome do mês por extenso
        nome_mes = calendar.month_name[mes_atual]

        # Formatar o dia com 2 dígitos
        dia_atual = hoje.day

        # Montar a string final
        data_formatada = f"{dia_atual:02d} de {nome_mes} de {ano_atual}"

        if not nome_documento:
            progress_label.config(text="Nome da empresa não encontrado.")
            print("Nome da empresa não encontrado.")
        else:
            # Ler o arquivo modelo e fazer as substituições
            template_doc = Document(template_editado)
            replacements = {
                'NOME DA EMPRESA': nome_documento,
                'JUNHO DE 2023': data_documento,
                '00.06.2023' : data_formatacao_documento,
                'XX.XXX.XXX/XXXX-XX': cnpj,
                '00/00/2000': data_documento_empresa,
                'DATA DA ABERTURA DA EMPRESA': converter_data_pt_br(infos_cartao_cnpj.get('data_abertura')),
                'cnpj': cnpj,
                'dataAbertura': format_date(infos_cartao_cnpj.get('data_abertura')),
                'nome_empresa': infos_cartao_cnpj.get('nome_empresa'),
                'nomeFantasia': infos_cartao_cnpj.get('nome_fantasia'),
                'porte': infos_cartao_cnpj.get('porte'),
                'codigoDescricao': infos_cartao_cnpj.get('codigo_completo'),
                'codigoDescSec': infos_cartao_cnpj.get('atividade_sec_text'),
                'codigo_desc_nat': "*****",
                'logradouro': infos_cartao_cnpj.get('logradouro'),
                'numero': infos_cartao_cnpj.get('numero'),
                'complemento': infos_cartao_cnpj.get('complemento'),
                'cep': infos_cartao_cnpj.get('cep'),
                'bairro': infos_cartao_cnpj.get('bairro'),
                'municipio': infos_cartao_cnpj.get('municipio'),
                'uf': infos_cartao_cnpj.get('uf'),
                'email': ', '.join(infos_cartao_cnpj.get('emails')),
                'telefone': ', '.join(infos_cartao_cnpj.get('telefones')),
                'situacao': infos_cartao_cnpj.get('status_text'),
                'dataSitCadastral': format_date(infos_cartao_cnpj.get('data_sit_cad')),
                'situacaoEspecial': "*****",
                'dataSituacaoEsp': "*****",
                'ENDEREÇO': infos_cartao_cnpj.get('logradouro')+', '+infos_cartao_cnpj.get('numero') +' - ' + infos_cartao_cnpj.get('bairro') + ' - ' + infos_cartao_cnpj.get('municipio') + ' - ' + infos_cartao_cnpj.get('uf'),
                '00 de maio de 2023': data_formatada
            }
            
            caminho_final_editado = output_pdf_path + '\\' + str(ano_atual) + ' - LTCAT - ' + nome_documento
            progress_label.config(text="Iniciando a substituição dos índices do documento.")
            substituir_texto_no_documento(template_doc, replacements, caminho_final_editado+'.docx', nome_documento, data_documento)

            # Salvar o novo documento modificado
            progress_label.config(text="Salvando o documento formado DOCX")
                        
            template_doc.save(output_docx_path)

            # Converter e salvar como PDF
            progress_label.config(text="Salvando o documento formado PDF")
            save_as_pdf(output_docx_path, caminho_final_editado+'.pdf')

            progress_label.config(text="Processo finalizado com sucesso!")
            progress_bar.stop()
            pythoncom.CoUninitialize()
            print("Documentos gerados com sucesso!")

#Função para iniciar a execução em uma thread separada
def start_process():
    progress_bar.start()
    threading.Thread(target=processar_arquivos, args=(
        progress_label, progress_bar)).start()
    
# Interface gráfica com Tkinter
root = tk.Tk()
root.title("Processar Arquivos LTCAT NR20")
root.geometry("400x300")

# Logo da empresa
logo_image = Image.open(fr"C:\Users\{USERNAME}\Desktop\ltcat\logo_empresa.jpg")
logo_image = logo_image.resize((200, 100), Image.LANCZOS)
logo_photo = ImageTk.PhotoImage(logo_image)
logo_label = tk.Label(root, image=logo_photo)
logo_label.pack(pady=10)

# Botão para processar arquivos
botao_processar = tk.Button(
    root, text="Processar arquivos LTCAT NR20", command=start_process)
botao_processar.pack(pady=10)

# Barra de progresso
progress_bar = ttk.Progressbar(
    root, orient="horizontal", mode="indeterminate", length=280)
progress_bar.pack(pady=10)

# Label de status do processo
progress_label = tk.Label(root, text="Aguardando...")
progress_label.pack()

# Iniciar a interface Tkinter
root.mainloop()
